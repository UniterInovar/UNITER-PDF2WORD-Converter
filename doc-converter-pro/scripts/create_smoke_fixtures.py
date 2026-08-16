from pathlib import Path

from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont

root = Path('/tmp/doc-converter-smoke')
root.mkdir(parents=True, exist_ok=True)

image_path = root / 'phenol-structure.png'
canvas = Image.new('RGB', (900, 420), 'white')
draw = ImageDraw.Draw(canvas)
font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 36)
small_font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 28)
center = (330, 215)
points = [(330, 92), (436, 154), (436, 276), (330, 338), (224, 276), (224, 154)]
for start, end in zip(points, points[1:] + points[:1]):
    draw.line([start, end], fill='#163d31', width=6)
inner = [(330, 116), (414, 164), (414, 266), (330, 314), (246, 266), (246, 164)]
for start, end in [(inner[0], inner[1]), (inner[2], inner[3]), (inner[4], inner[5])]:
    draw.line([start, end], fill='#163d31', width=4)
draw.line([(436, 154), (555, 86)], fill='#163d31', width=6)
draw.text((570, 56), 'OH', fill='#163d31', font=font)
draw.text((272, 168), '⌬', fill='#5b9a72', font=font)
draw.text((55, 168), 'Phenol  C₆H₅OH', fill='#315948', font=small_font)
canvas.save(image_path)

pdf_path = root / 'chemistry-layout.pdf'
pdf = FPDF()
pdf.add_page()
pdf.set_font('Helvetica', 'B', 18)
pdf.cell(0, 12, 'Chemistry layout smoke test')
pdf.ln(14)
pdf.set_font('Helvetica', '', 11)
pdf.multi_cell(0, 7, 'Text layout, an embedded structural formula, and table content should remain readable after conversion.')
pdf.ln(4)
pdf.set_font('Helvetica', 'B', 12)
pdf.cell(0, 8, 'Formula label: C6H5OH')
pdf.ln(12)
pdf.image(str(image_path), x=22, w=166)
pdf.ln(84)
pdf.set_font('Helvetica', '', 10)
for label, value in [('Compound', 'Phenol'), ('Atoms', 'C6 H6 O1'), ('Geometry', 'Aromatic ring with hydroxyl group')]:
    pdf.cell(48, 8, label, border=1)
    pdf.cell(130, 8, value, border=1)
    pdf.ln(8)
pdf.output(str(pdf_path))

docx_path = root / 'technical-layout.docx'
word = Document()
word.add_heading('Technical layout smoke test', level=1)
word.add_paragraph('This source tests a Word to PDF path with a structural formula image, tables, and technical text.')
word.add_paragraph('Formula label: C6H5OH')
word.add_picture(str(image_path), width=Inches(5.7))
table = word.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Element'
table.rows[0].cells[1].text = 'Description'
for element, description in [('Image', 'Embedded chemical structural formula'), ('Table', 'Structured rows and columns'), ('Formula', 'Phenol / C6H5OH')]:
    cells = table.add_row().cells
    cells[0].text = element
    cells[1].text = description
word.save(docx_path)
print(pdf_path)
print(docx_path)
